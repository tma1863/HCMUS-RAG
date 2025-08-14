from typing import List
from docx import Document

class TableExtractor:
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.tables: List[List[List[str]]] = []

    def extract_tables_with_headings(self):
        paragraphs = list(self.doc.paragraphs)
        for table in self.doc.tables:
            prev_paragraph_text = ""
            for para in paragraphs:
                if table._element is para._element.getnext():
                    prev_paragraph_text = para.text.strip()
                    break
            course_info = ["Course id", prev_paragraph_text]
            table_data = [course_info]
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            self.tables.append(table_data)
        return self.tables